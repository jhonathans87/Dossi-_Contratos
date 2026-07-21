
import io
import os
import re
import json
import hashlib
from datetime import date, datetime
from typing import List, Dict, Any, Tuple

import streamlit as st
from pypdf import PdfReader
from docx import Document
from openai import OpenAI
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
)

APP_TITLE = "Dossiê Executivo Imobiliário"

MASTER_PROMPT = """
Você é um analista financeiro, patrimonial e contratual especializado em unidades imobiliárias.

Crie um DOSSIÊ EXECUTIVO COMPLETO DE VALORIZAÇÃO PATRIMONIAL E ANÁLISE CONTRATUAL usando exclusivamente:
1. os documentos fornecidos;
2. o valor atual de tabela informado;
3. a data-base informada.

REGRAS OBRIGATÓRIAS
- Não presuma informações.
- Não invente cláusulas, páginas, taxas, índices ou datas.
- Quando algo não estiver disponível, escreva: "Documento não localizado entre os arquivos analisados." ou "Informação não localizada nos documentos analisados.", conforme o caso.
- Aplique a hierarquia: aditivo mais recente; aditivos anteriores; contrato original; extrato financeiro; informação do usuário.
- O aditivo prevalece apenas sobre o que alterar expressamente.
- Sempre diferencie valor original, valor ajustado, valor corrigido e recebimento líquido.
- Não apresente valorização bruta como lucro líquido.
- Não omita saldo devedor.
- Não afirme superioridade sobre o CDI sem cálculo exato.
- Patrimônio de afetação somente pode ser afirmado se constar expressamente nos documentos.
- Em temas jurídicos, indique documento, página e cláusula/item quando localizados.
- Para cada dado relevante, indique a origem.
- Destaque divergências entre contrato, aditivos e extrato.
- Não trate o último número de "Par" como quantidade de parcelas pagas; conte os registros por tipo.
- Caso a série histórica oficial do CDI não tenha sido fornecida ao sistema, escreva exatamente:
"Simulação de CDI não concluída por ausência da série histórica oficial necessária à capitalização individual dos aportes. Nenhum valor estimado foi inserido."

ESTRUTURA OBRIGATÓRIA
1. Identificação da unidade
2. Resumo executivo
3. Documentos analisados e documentos ausentes
4. Hierarquia documental aplicada
5. Dados contratuais - condição original
6. Dados contratuais - condição após aditivos
7. Alterações efetivas
8. Conferência financeira
9. Inconsistências identificadas
10. Evolução patrimonial
11. Rentabilidade sobre o contrato
12. Rentabilidade sobre o capital investido
13. Multiplicador patrimonial
14. Patrimônio líquido estimado
15. Simulação em 100% do CDI
16. Comparativo imóvel x CDI
17. Prazo de entrega
18. Multa por atraso da incorporadora
19. Inadimplência do adquirente
20. Patrimônio de afetação
21. Quadro executivo de rentabilidade
22. Ranking de performance
23. Conclusão executiva

FÓRMULAS
- Valorização sobre o contrato original = valor atual de tabela - valor original do contrato.
- Valorização sobre valor ajustado = valor atual de tabela - valor contratual ajustado.
- Rentabilidade sobre o contrato = valorização / valor contratual de referência x 100.
- Rentabilidade sobre o capital investido = valorização / recebimento líquido x 100.
- Multiplicador patrimonial = valor atual de tabela / recebimento líquido.
- Patrimônio líquido estimado = valor atual de tabela - saldo devedor total atualizado.
- Ganho econômico líquido estimado = patrimônio líquido estimado - recebimento líquido.

CLASSIFICAÇÃO
- Resultado negativo: Baixo Desempenho
- Até 30%: Regular
- Acima de 30% até 60%: Bom
- Acima de 60% até 100%: Muito Bom
- Acima de 100%: Excelente

FORMATAÇÃO
- Escreva em português do Brasil.
- Use títulos e subtítulos claros.
- Use tabelas em Markdown quando isso melhorar a leitura.
- Linguagem técnica, persuasiva, responsável e adequada para diretoria, cliente, jurídico, comercial e financeiro.
"""

def clean_text(text: str) -> str:
    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()

def extract_pdf(file_bytes: bytes) -> Tuple[str, int]:
    reader = PdfReader(io.BytesIO(file_bytes))
    pages = []
    for idx, page in enumerate(reader.pages, start=1):
        page_text = page.extract_text() or ""
        pages.append(f"\n--- PÁGINA {idx} ---\n{page_text}")
    return clean_text("\n".join(pages)), len(reader.pages)

def extract_docx(file_bytes: bytes) -> str:
    doc = Document(io.BytesIO(file_bytes))
    parts = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text.strip() for cell in row.cells))
    return clean_text("\n".join(parts))

def extract_txt(file_bytes: bytes) -> str:
    for enc in ("utf-8", "latin-1"):
        try:
            return clean_text(file_bytes.decode(enc))
        except UnicodeDecodeError:
            pass
    return ""

def extract_uploaded_file(uploaded) -> Dict[str, Any]:
    raw = uploaded.getvalue()
    suffix = uploaded.name.lower().rsplit(".", 1)[-1]
    if suffix == "pdf":
        text, pages = extract_pdf(raw)
    elif suffix == "docx":
        text, pages = extract_docx(raw), None
    elif suffix in ("txt", "md", "csv"):
        text, pages = extract_txt(raw), None
    else:
        text, pages = "", None
    return {
        "name": uploaded.name,
        "size": len(raw),
        "pages": pages,
        "text": text,
        "sha256": hashlib.sha256(raw).hexdigest()
    }

def truncate_documents(docs: List[Dict[str, Any]], limit_chars: int = 220000) -> str:
    blocks = []
    used = 0
    for i, doc in enumerate(docs, start=1):
        header = (
            f"\n\n===== DOCUMENTO {i}: {doc['name']} =====\n"
            f"Páginas: {doc.get('pages') or 'não aplicável'}\n"
            f"SHA-256: {doc['sha256']}\n"
        )
        available = max(0, limit_chars - used - len(header))
        if available <= 0:
            break
        body = doc["text"][:available]
        blocks.append(header + body)
        used += len(header) + len(body)
    return "".join(blocks)

def call_ai(documents_text: str, table_value: float, base_date: date, notes: str) -> str:
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        raise RuntimeError(
            "A variável OPENAI_API_KEY não foi configurada no ambiente de hospedagem."
        )
    model = os.getenv("OPENAI_MODEL", "gpt-5-mini")
    client = OpenAI(api_key=api_key)

    user_input = f"""
VALOR ATUAL DE TABELA: R$ {table_value:,.2f}
DATA-BASE: {base_date.strftime('%d/%m/%Y')}
OBSERVAÇÕES DO USUÁRIO:
{notes or 'Nenhuma observação adicional.'}

DOCUMENTOS EXTRAÍDOS:
{documents_text}
"""
    response = client.responses.create(
        model=model,
        instructions=MASTER_PROMPT,
        input=user_input,
    )
    return response.output_text

def make_pdf(report_text: str, title: str) -> bytes:
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=42,
        leftMargin=42,
        topMargin=48,
        bottomMargin=48,
        title=title,
    )
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(
        name="TitleCenter",
        parent=styles["Title"],
        alignment=TA_CENTER,
        spaceAfter=18,
        fontSize=18,
        leading=22,
    ))
    styles.add(ParagraphStyle(
        name="BodyCustom",
        parent=styles["BodyText"],
        fontSize=9.2,
        leading=13,
        spaceAfter=6,
    ))
    styles.add(ParagraphStyle(
        name="HeadingCustom",
        parent=styles["Heading2"],
        fontSize=12,
        leading=15,
        spaceBefore=10,
        spaceAfter=6,
    ))

    story = [Paragraph(title, styles["TitleCenter"]), Spacer(1, 8)]
    lines = report_text.splitlines()
    table_rows = []
    in_table = False

    def flush_table():
        nonlocal table_rows, in_table
        if table_rows:
            normalized = []
            max_cols = max(len(r) for r in table_rows)
            for row in table_rows:
                normalized.append(row + [""] * (max_cols - len(row)))
            table = Table(normalized, repeatRows=1, hAlign="LEFT")
            table.setStyle(TableStyle([
                ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
                ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 7.5),
                ("LEADING", (0, 0), (-1, -1), 9.5),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ]))
            story.extend([table, Spacer(1, 8)])
        table_rows = []
        in_table = False

    for raw_line in lines:
        line = raw_line.strip()
        if line.startswith("|") and line.endswith("|"):
            cells = [c.strip() for c in line.strip("|").split("|")]
            if all(re.fullmatch(r":?-{3,}:?", c or "") for c in cells):
                continue
            table_rows.append([Paragraph(c or " ", styles["BodyCustom"]) for c in cells])
            in_table = True
            continue
        if in_table:
            flush_table()
        if not line:
            story.append(Spacer(1, 5))
        elif line.startswith("### "):
            story.append(Paragraph(line[4:], styles["HeadingCustom"]))
        elif line.startswith("## "):
            story.append(Paragraph(line[3:], styles["HeadingCustom"]))
        elif line.startswith("# "):
            story.append(Paragraph(line[2:], styles["HeadingCustom"]))
        elif re.match(r"^\d+\.\s+", line):
            story.append(Paragraph(f"<b>{line}</b>", styles["HeadingCustom"]))
        elif line.startswith("- "):
            story.append(Paragraph("• " + line[2:], styles["BodyCustom"]))
        else:
            safe = (
                line.replace("&", "&amp;")
                    .replace("<", "&lt;")
                    .replace(">", "&gt;")
            )
            safe = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", safe)
            story.append(Paragraph(safe, styles["BodyCustom"]))
    flush_table()
    doc.build(story)
    return buffer.getvalue()

st.set_page_config(page_title=APP_TITLE, page_icon="🏢", layout="wide")

st.title("🏢 Dossiê Executivo Imobiliário")
st.caption(
    "Anexe contrato, aditivos e extrato; informe o valor atual de tabela; revise os dados e gere o dossiê."
)

with st.sidebar:
    st.header("Configuração")
    st.info(
        "A aplicação utiliza IA para interpretar os documentos. "
        "Os cálculos e conclusões devem ser revisados antes do uso jurídico ou financeiro."
    )
    st.text_input(
        "Modelo de IA",
        value=os.getenv("OPENAI_MODEL", "gpt-5-mini"),
        disabled=True
    )

col1, col2 = st.columns([2, 1])

with col1:
    uploads = st.file_uploader(
        "Documentos da unidade",
        type=["pdf", "docx", "txt", "md", "csv"],
        accept_multiple_files=True,
        help="Envie contrato original, todos os aditivos, extrato e documentos registrais disponíveis."
    )

with col2:
    table_value = st.number_input(
        "Valor atual de tabela (R$)",
        min_value=0.0,
        value=0.0,
        step=1000.0,
        format="%.2f"
    )
    base_date = st.date_input("Data-base", value=date.today())

notes = st.text_area(
    "Observações complementares",
    placeholder="Ex.: valor informado pelo comercial, particularidade da negociação, documento ausente..."
)

if uploads:
    st.subheader("Arquivos selecionados")
    preview_rows = []
    for f in uploads:
        preview_rows.append({
            "Arquivo": f.name,
            "Tamanho (KB)": round(f.size / 1024, 1),
            "Tipo": f.type or "não identificado"
        })
    st.dataframe(preview_rows, use_container_width=True, hide_index=True)

process = st.button(
    "Processar e gerar dossiê",
    type="primary",
    use_container_width=True,
    disabled=not uploads or table_value <= 0
)

if process:
    with st.status("Processando documentos...", expanded=True) as status:
        try:
            docs = []
            for uploaded in uploads:
                st.write(f"Extraindo: {uploaded.name}")
                item = extract_uploaded_file(uploaded)
                if not item["text"]:
                    st.warning(
                        f"Não foi possível extrair texto de {uploaded.name}. "
                        "PDFs digitalizados podem exigir OCR."
                    )
                docs.append(item)

            usable_docs = [d for d in docs if d["text"]]
            if not usable_docs:
                raise RuntimeError("Nenhum texto legível foi extraído dos documentos.")

            documents_text = truncate_documents(usable_docs)
            st.write("Analisando conteúdo contratual e financeiro...")
            report = call_ai(
                documents_text=documents_text,
                table_value=table_value,
                base_date=base_date,
                notes=notes
            )
            st.session_state["report"] = report
            st.session_state["report_meta"] = {
                "table_value": table_value,
                "base_date": base_date.isoformat(),
                "files": [d["name"] for d in docs],
            }
            status.update(label="Dossiê concluído", state="complete", expanded=False)
        except Exception as exc:
            status.update(label="Falha no processamento", state="error", expanded=True)
            st.error(str(exc))

if "report" in st.session_state:
    report = st.session_state["report"]
    st.divider()
    st.subheader("Dossiê gerado")
    st.markdown(report)

    pdf_bytes = make_pdf(
        report,
        "Dossiê Executivo de Valorização Patrimonial e Análise Contratual"
    )
    col_a, col_b = st.columns(2)
    with col_a:
        st.download_button(
            "Baixar dossiê em PDF",
            data=pdf_bytes,
            file_name="dossie_executivo_imobiliario.pdf",
            mime="application/pdf",
            use_container_width=True
        )
    with col_b:
        st.download_button(
            "Baixar memória em Markdown",
            data=report.encode("utf-8"),
            file_name="dossie_executivo_imobiliario.md",
            mime="text/markdown",
            use_container_width=True
        )

    st.warning(
        "Antes de apresentar o documento a cliente, diretoria, jurídico ou financeiro, "
        "revise páginas, cláusulas, valores, datas e eventuais falhas de extração."
    )
